#!/usr/bin/env python
"""Convert EMARX Markdown papers into explicitly styled academic DOCX files."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


CITATION_RE = re.compile(r"(?<!!)\[(\d+)\]")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
REFERENCE_RE = re.compile(r"^\s*\[(\d+)\]\s*(.+?)\s*$")
ABSTRACT_RE = re.compile(r"^\s*(摘要)\s*[：:]\s*(.*)$")
KEYWORDS_RE = re.compile(r"^\s*(关键词|关键字)\s*[：:]\s*(.*)$")

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"


def clean_inline_markdown(text: str) -> str:
    text = text.replace("\ufeff", "")
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def set_run_font(
    run,
    size: float,
    *,
    east_asia: str = BODY_FONT,
    bold: bool = False,
    superscript: bool = False,
) -> None:
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.superscript = superscript
    run.font.color.rgb = RGBColor(0, 0, 0)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), east_asia)
    color = run._element.get_or_add_rPr().get_or_add_color()
    color.attrib.pop(qn("w:themeColor"), None)
    color.attrib.pop(qn("w:themeTint"), None)
    color.attrib.pop(qn("w:themeShade"), None)


def configure_style(
    style,
    *,
    size: float,
    east_asia: str = BODY_FONT,
    bold: bool = False,
    alignment=None,
    first_line: float | None = None,
    line_spacing: float = 1.5,
    before: float = 0,
    after: float = 0,
    keep_with_next: bool = False,
) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), east_asia)
    color = style._element.get_or_add_rPr().get_or_add_color()
    color.attrib.pop(qn("w:themeColor"), None)
    color.attrib.pop(qn("w:themeTint"), None)
    color.attrib.pop(qn("w:themeShade"), None)
    fmt = style.paragraph_format
    if alignment is not None:
        fmt.alignment = alignment
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Pt(first_line) if first_line is not None else None
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def ensure_style(document: Document, name: str):
    styles = document.styles
    return styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = 0
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)

    configure_style(
        document.styles["Normal"],
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line=24,
    )
    configure_style(
        ensure_style(document, "EMARX Title"),
        size=18,
        east_asia=HEADING_FONT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.25,
        after=12,
        keep_with_next=True,
    )
    configure_style(
        document.styles["Heading 1"],
        size=14,
        east_asia=HEADING_FONT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.25,
        before=12,
        after=6,
        keep_with_next=True,
    )
    configure_style(
        document.styles["Heading 2"],
        size=12,
        east_asia=HEADING_FONT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.25,
        before=9,
        after=3,
        keep_with_next=True,
    )
    configure_style(
        document.styles["Heading 3"],
        size=12,
        east_asia=BODY_FONT,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.25,
        before=6,
        after=3,
        keep_with_next=True,
    )
    configure_style(
        ensure_style(document, "EMARX Abstract"),
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.25,
        after=3,
    )
    configure_style(
        ensure_style(document, "EMARX Keywords"),
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.25,
        after=9,
    )
    configure_style(
        ensure_style(document, "EMARX Reference"),
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.25,
        after=2,
    )
    ref_fmt = document.styles["EMARX Reference"].paragraph_format
    ref_fmt.left_indent = Pt(21)
    ref_fmt.first_line_indent = Pt(-21)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def add_inline_runs(paragraph, text: str, *, size: float = 12, citations: bool = True) -> int:
    text = clean_inline_markdown(text)
    cursor = 0
    superscript_count = 0
    matches = list(CITATION_RE.finditer(text)) if citations else []
    for match in matches:
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size)
        run = paragraph.add_run(match.group(0))
        set_run_font(run, max(8, size - 2), superscript=True)
        superscript_count += 1
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size)
    return superscript_count


def add_labeled_paragraph(document: Document, label: str, text: str, style: str) -> int:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(f"{label}：")
    set_run_font(run, 10.5, east_asia=HEADING_FONT, bold=True)
    return add_inline_runs(paragraph, text, size=10.5)


def tokenize(markdown: str) -> list[tuple[str, str, int | None]]:
    """Return semantic blocks while joining hard-wrapped lines within paragraphs."""
    blocks: list[tuple[str, str, int | None]] = []
    paragraph_lines: list[str] = []
    in_code = False

    def flush() -> None:
        if paragraph_lines:
            blocks.append(("paragraph", "".join(line.strip() for line in paragraph_lines), None))
            paragraph_lines.clear()

    for raw in markdown.replace("\r\n", "\n").split("\n"):
        line = raw.rstrip().replace("\ufeff", "")
        if line.strip().startswith("```"):
            flush()
            in_code = not in_code
            continue
        if in_code:
            paragraph_lines.append(line)
            continue
        if not line.strip():
            flush()
            continue
        heading = HEADING_RE.match(line)
        if heading:
            flush()
            blocks.append(("heading", clean_inline_markdown(heading.group(2)), len(heading.group(1))))
            continue
        reference = REFERENCE_RE.match(line)
        if reference:
            flush()
            blocks.append(("reference", f"[{reference.group(1)}] {reference.group(2)}", None))
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            flush()
            blocks.append(("paragraph", (bullet or numbered).group(1), None))
            continue
        paragraph_lines.append(line)
    flush()
    return blocks


def convert(markdown: str, output: Path) -> dict:
    document = Document()
    configure_document(document)
    add_page_number(document.sections[0])
    document.core_properties.title = "EMARX 学术论文"
    document.core_properties.subject = "中文学理思辨论文"

    title_written = False
    in_references = False
    stats = {
        "title_count": 0,
        "heading_counts": {"1": 0, "2": 0, "3": 0},
        "body_paragraphs": 0,
        "abstract_count": 0,
        "keywords_count": 0,
        "reference_count": 0,
        "superscript_citation_count": 0,
    }
    pending_labeled_block: str | None = None

    for block_type, raw_text, level in tokenize(markdown):
        text = clean_inline_markdown(raw_text)
        if not text:
            continue

        if block_type == "heading":
            if level == 1 and not title_written:
                paragraph = document.add_paragraph(style="EMARX Title")
                run = paragraph.add_run(text)
                set_run_font(run, 18, east_asia=HEADING_FONT, bold=True)
                title_written = True
                stats["title_count"] += 1
                continue
            if text in {"摘要", "关键词", "关键字"}:
                pending_labeled_block = "关键词" if text in {"关键词", "关键字"} else "摘要"
                continue
            if text == "参考文献":
                in_references = True
            style_level = 1 if level is None or level <= 2 else min(level - 1, 3)
            paragraph = document.add_paragraph(style=f"Heading {style_level}")
            run = paragraph.add_run(text)
            set_run_font(
                run,
                14 if style_level == 1 else 12,
                east_asia=HEADING_FONT if style_level < 3 else BODY_FONT,
                bold=True,
            )
            stats["heading_counts"][str(style_level)] += 1
            continue

        if block_type == "reference" or (in_references and REFERENCE_RE.match(text)):
            paragraph = document.add_paragraph(style="EMARX Reference")
            add_inline_runs(paragraph, text, size=10.5, citations=False)
            stats["reference_count"] += 1
            continue

        abstract = ABSTRACT_RE.match(text)
        if abstract:
            stats["superscript_citation_count"] += add_labeled_paragraph(
                document, abstract.group(1), abstract.group(2), "EMARX Abstract"
            )
            stats["abstract_count"] += 1
            continue
        keywords = KEYWORDS_RE.match(text)
        if keywords:
            stats["superscript_citation_count"] += add_labeled_paragraph(
                document, "关键词", keywords.group(2), "EMARX Keywords"
            )
            stats["keywords_count"] += 1
            continue

        if pending_labeled_block == "摘要":
            stats["superscript_citation_count"] += add_labeled_paragraph(
                document, "摘要", text, "EMARX Abstract"
            )
            stats["abstract_count"] += 1
            pending_labeled_block = None
            continue

        if pending_labeled_block == "关键词":
            stats["superscript_citation_count"] += add_labeled_paragraph(
                document, "关键词", text, "EMARX Keywords"
            )
            stats["keywords_count"] += 1
            pending_labeled_block = None
            continue

        paragraph = document.add_paragraph(style="Normal")
        stats["superscript_citation_count"] += add_inline_runs(paragraph, text)
        stats["body_paragraphs"] += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    stats["output"] = str(output.resolve())
    stats["output_bytes"] = output.stat().st_size
    return stats


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="UTF-8 Markdown file")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--audit-output", help="Optional JSON conversion audit")
    args = parser.parse_args()

    source = Path(args.input)
    target = Path(args.output)
    stats = convert(source.read_text(encoding="utf-8"), target)
    if args.audit_output:
        audit_path = Path(args.audit_output)
        audit_path.parent.mkdir(parents=True, exist_ok=True)
        audit_path.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(stats, ensure_ascii=False))


if __name__ == "__main__":
    main()
